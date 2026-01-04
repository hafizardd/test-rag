import os
import json
import time
import random
from pathlib import Path
from typing import List, Dict, Tuple, Any, Optional
from dataclasses import dataclass
from datetime import datetime
from tqdm import tqdm
from collections import Counter

import numpy as np
from groq import Groq
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from supabase import create_client, Client

# Import PDF libraries with fallback
UNSTRUCTURED_AVAILABLE = False
try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition. docx import partition_docx
    UNSTRUCTURED_AVAILABLE = True
    print("✅ Unstructured library available")
except ImportError: 
    print("⚠️ Unstructured library not available, using fallback PDF processing")

# Fallback PDF libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("⚠️ pdfplumber not available")

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("⚠️ PyPDF2 not available")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("⚠️ python-docx not available")

# Load environment variables
PROJECT_ROOT = Path(__file__).resolve().parent.parent
ENV_PATH = PROJECT_ROOT / ".env"
load_dotenv(dotenv_path=ENV_PATH)


# ========== Configuration ==========
def get_groq_client() -> Groq:
    """Get GROQ API client"""
    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("GROQ_API_KEY not set.  Provide it via env or a .env file.")
    return Groq(api_key=api_key)


def get_supabase_client() -> Client:
    """Get Supabase client"""
    url = os. environ.get("SUPABASE_URL", "").strip()
    key = os.environ.get("SUPABASE_SERVICE_KEY", "").strip()
    
    if not url or not key: 
        raise RuntimeError(
            "SUPABASE_URL and SUPABASE_SERVICE_KEY must be set in .env file.\n"
            "Get these from:  Supabase Dashboard → Settings → API"
        )
    
    return create_client(url, key)


# ========== Document Metadata Dataclass ==========
@dataclass
class DocumentMetadata:
    """Metadata untuk dokumen"""
    filename: str
    file_size: int
    creation_date: datetime
    page_count: int = 0
    keywords: Optional[List[str]] = None
    summary: str = ""
    document_type:  str = "general"
    doc_id: str = ""


@dataclass
class DocChunk:
    """Document chunk with metadata"""
    doc_id:  str
    chunk_id: int
    text: str
    meta: Dict
    embedding_model: str = "sentence-transformers/all-mpnet-base-v2"
    chunk_size: int = 1000
    chunk_overlap: int = 100
    metadata: Optional[DocumentMetadata] = None
    content_type: str = "text"
    page_number: Optional[int] = None


# ========== Simple Element Class for Fallback ==========
class SimpleElement: 
    """Simple element class for fallback PDF processing"""
    def __init__(self, text: str, page_num: Optional[int] = None, element_type: str = "text"):
        self.text = text
        self.element_type = element_type
        self.metadata = type('obj', (object,), {
            'page_number': page_num,
            'text_as_html': text if element_type == "table" else None
        })()
    
    def __str__(self):
        return self.text


# ========== Multimodal Extraction Functions ==========
def extract_elements_from_pdf(file_path: str, output_path: str = "./content/") -> Tuple[List, List[str], Dict]:
    """Extract elements from PDF including text, tables, and images with fallback support"""
    print(f"🔍 Extracting from PDF: {file_path}")
    
    if UNSTRUCTURED_AVAILABLE: 
        try:
            print("📄 Trying unstructured library...")
            chunks = partition_pdf(
                filename=file_path,
                infer_table_structure=True,
                strategy="hi_res",
                extract_image_block_types=["Image"],
                extract_image_block_to_payload=True,
                chunking_strategy="by_title",
                max_characters=10000,
                combine_text_under_n_chars=2000,
                new_after_n_chars=6000,
                languages=["eng", "ind"]
            )
            
            images_base64 = []
            page_map = {}
            
            for idx, chunk in enumerate(chunks):
                if hasattr(chunk, 'metadata') and hasattr(chunk.metadata, 'page_number'):
                    page_map[idx] = chunk.metadata.page_number
                
                if "CompositeElement" in str(type(chunk)):
                    chunk_els = chunk.metadata.orig_elements
                    for el in chunk_els:
                        if "Image" in str(type(el)):
                            if hasattr(el. metadata, 'image_base64') and el.metadata.image_base64:
                                images_base64.append(el.metadata. image_base64)
            
            print(f"✅ Extracted {len(chunks)} chunks with unstructured")
            return chunks, images_base64, page_map
        
        except Exception as e:
            print(f"⚠️ Unstructured failed: {str(e)}")
    
    # Method 2: Fallback with pdfplumber
    if PDFPLUMBER_AVAILABLE:
        try:
            elements = []
            images_base64 = []
            page_map = {}
            
            with pdfplumber.open(file_path) as pdf:
                print(f"📖 Processing {len(pdf.pages)} pages with pdfplumber...")
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text and text.strip():
                        element = SimpleElement(text, page_num, "text")
                        elements.append(element)
                        page_map[len(elements) - 1] = page_num
                    
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table: 
                                table_text = "\n".join([" | ".join([str(cell) if cell else "" for cell in row]) for row in table])
                                table_html = f"<table>{table_text}</table>"
                                table_element = SimpleElement(table_html, page_num, "table")
                                elements.append(table_element)
                                page_map[len(elements) - 1] = page_num
            
            print(f"✅ Extracted {len(elements)} elements with pdfplumber")
            return elements, images_base64, page_map
        
        except Exception as e:
            print(f"⚠️ pdfplumber failed: {str(e)}")
    
    # Method 3: Last fallback with PyPDF2
    if PYPDF2_AVAILABLE: 
        try:
            print("📄 Using PyPDF2 as last fallback...")
            elements = []
            images_base64 = []
            page_map = {}
            
            reader = PdfReader(file_path)
            
            for page_num, page in enumerate(reader.pages, start=1):
                text = page. extract_text()
                if text and text.strip():
                    element = SimpleElement(text, page_num, "text")
                    elements.append(element)
                    page_map[len(elements) - 1] = page_num
            
            print(f"✅ Extracted {len(elements)} elements with PyPDF2")
            return elements, images_base64, page_map
        
        except Exception as e:
            print(f"❌ PyPDF2 failed: {str(e)}")
    
    print(f"❌ All PDF extraction methods failed for {file_path}")
    return [], [], {}


def extract_elements_from_docx(file_path: str) -> Tuple[List, List[str], Dict]:
    """Extract elements from DOCX with fallback support"""
    print(f"🔍 Extracting from DOCX: {file_path}")
    
    if UNSTRUCTURED_AVAILABLE:
        try:
            print("📄 Trying unstructured library...")
            chunks = partition_docx(
                filename=file_path,
                chunking_strategy="by_title",
                max_characters=10000,
                combine_text_under_n_chars=2000,
                new_after_n_chars=6000,
            )
            
            images_base64 = []
            page_map = {}
            
            for idx, chunk in enumerate(chunks):
                if hasattr(chunk, 'metadata') and hasattr(chunk.metadata, 'page_number'):
                    page_map[idx] = chunk.metadata. page_number
            
            print(f"✅ Extracted {len(chunks)} chunks with unstructured")
            return chunks, images_base64, page_map
        
        except Exception as e:
            print(f"⚠️ Unstructured failed: {str(e)}")
    
    if PYTHON_DOCX_AVAILABLE: 
        try:
            print("📄 Using python-docx fallback...")
            doc = Document(file_path)
            elements = []
            page_map = {}
            
            for para in doc.paragraphs:
                if para.text.strip():
                    element = SimpleElement(para.text, None, "text")
                    elements.append(element)
            
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row. cells])
                    table_text += row_text + "\n"
                
                if table_text.strip():
                    table_element = SimpleElement(f"<table>{table_text}</table>", None, "table")
                    elements.append(table_element)
            
            print(f"✅ Extracted {len(elements)} elements with python-docx")
            return elements, [], page_map
        
        except Exception as e:
            print(f"❌ python-docx failed: {str(e)}")
    
    print(f"❌ All DOCX extraction methods failed for {file_path}")
    return [], [], {}


# ========== Summarization Functions ==========
def summarize_image_with_groq(image_base64: str, groq_client: Groq) -> str:
    """Summarize image using Groq's Llama 4 Scout Vision model"""
    try:
        prompt = """Describe this image in detail.  
Focus on the key visual elements, text content, charts, diagrams, or any important information. 
Be specific and concise."""
        
        response = groq_client.chat.completions. create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text":  prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}
                        }
                    ]
                }
            ],
            temperature=1,
            max_tokens=500,
            top_p=1,
            stream=False,
        )
        
        return response.choices[0].message. content
    except Exception as e:
        print(f"Error summarizing image with Groq: {str(e)}")
        return "Image content could not be processed."


def summarize_text_with_groq(text: str, groq_client: Groq, max_retries: int = 3) -> str:
    """Summarize text using Groq with auto-retry for connection errors."""
    if len(text. strip()) < 50:
        return text

    for attempt in range(max_retries):
        try:
            prompt = f"""You are an assistant tasked with summarizing text content. 
Give a concise summary of the following text. 
Respond only with the summary, no additional comment. 

Text:  {text}"""
            
            response = groq_client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[{"role": "user", "content":  prompt}],
                temperature=0.5,
                max_tokens=300
            )
            
            return response.choices[0].message.content

        except Exception as e:
            error_msg = str(e)
            print(f"⚠️ Attempt {attempt + 1}/{max_retries} failed: {error_msg}")
            
            if "Connection error" in error_msg or "429" in error_msg:
                sleep_time = (2 ** attempt) + random.uniform(0, 1)
                time.sleep(sleep_time)
            else:
                break

    print("❌ Failed to summarize after retries.  Using fallback.")
    return text[:500]


def summarize_table_with_groq(table_html: str, groq_client:  Groq) -> str:
    """Summarize table using Groq"""
    try:
        prompt = f"""You are an assistant tasked with summarizing tables.
Give a concise summary of the following table. 
Respond only with the summary, no additional comment.

Table: {table_html}"""
        
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=300
        )
        
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error summarizing table: {str(e)}")
        return table_html[:500]


# ========== Advanced Chunking Strategies ==========
def chunk_text(
    text: str,
    strategy: str = "recursive",
    max_tokens: int = 1000,
    overlap_tokens: int = 100,
) -> List[str]:
    """Advanced chunking dengan berbagai strategi"""
    if not text: 
        return []

    if strategy == "recursive": 
        text_splitter = RecursiveCharacterTextSplitter(
            separators=["\n\n", "\n", ". ", " ", ""],
            chunk_size=max_tokens,
            chunk_overlap=overlap_tokens,
            length_function=len,
        )
        return text_splitter.split_text(text)

    elif strategy == "semantic":
        sentences = text.split(". ")
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk + sentence) < max_tokens:
                current_chunk += sentence + ". "
            else:
                if current_chunk: 
                    chunks.append(current_chunk. strip())
                current_chunk = sentence + ". "

        if current_chunk: 
            chunks.append(current_chunk.strip())
        return chunks

    elif strategy == "paragraph":
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk + para) < max_tokens:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk. strip())
                current_chunk = para + "\n\n"

        if current_chunk:
            chunks.append(current_chunk. strip())
        return chunks

    else:  # simple
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0

        for word in words:
            word_length = len(word) + 1
            if current_length + word_length > max_tokens and current_chunk:
                chunks. append(" ".join(current_chunk))
                overlap_words = (
                    current_chunk[-overlap_tokens // 10:]
                    if len(current_chunk) > overlap_tokens // 10
                    else []
                )
                current_chunk = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk)
            else:
                current_chunk.append(word)
                current_length += word_length

        if current_chunk: 
            chunks.append(" ".join(current_chunk))

        return chunks


# ========== Document Metadata Helpers ==========
def extract_keywords_simple(text: str, top_n: int = 10) -> List[str]:
    """Extract simple keywords from text based on word frequency"""
    words = text.lower().split()
    stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were'}
    filtered_words = [w for w in words if w not in stopwords and len(w) > 3]
    
    word_counts = Counter(filtered_words)
    return [word for word, count in word_counts.most_common(top_n)]


def generate_document_summary(text: str, max_length: int = 200) -> str:
    """Generate simple document summary from first sentences"""
    sentences = text.split('.  ')
    summary = ''
    for sentence in sentences[: 3]: 
        if len(summary + sentence) < max_length:
            summary += sentence + '. '
        else:
            break
    return summary. strip() or text[: max_length] + "..."


# ========== Supabase Vector Store Class ==========
class SupabaseVectorStore:
    """Vector store using Supabase pgvector"""
    
    def __init__(self, embed_model: SentenceTransformer, table_name: str = "document_chunks"):
        self.client = get_supabase_client()
        self.embed_model = embed_model
        self.table_name = table_name
        self.dimension = 768  # for all-mpnet-base-v2
        
        # Verify connection
        try:
            count = self.get_document_count()
            print(f"✅ Connected to Supabase.  Current documents: {count}")
        except Exception as e:
            print(f"⚠️ Supabase connection warning: {str(e)}")
    
    def add_documents(
        self,
        documents: List[Dict[str, Any]],
        show_progress: bool = True
    ) -> int:
        """Add documents to Supabase vector store"""
        if not documents:
            return 0
        
        total_added = 0
        
        for doc in tqdm(documents, desc="Adding to Supabase", disable=not show_progress):
            try:
                # Generate embedding for this document
                text = doc["text"]
                embedding = self. embed_model.encode([text], convert_to_numpy=True)[0]
                embedding_list = [float(x) for x in embedding. tolist()]
                
                # Prepare metadata
                metadata = {
                    "category": doc. get("category", "unknown"),
                    "keywords": doc.get("keywords", ""),
                    "page_number": doc.get("page_number"),
                    "filename": doc.get("filename", doc.get("doc_id", "")),
                    "chunk_size": doc.get("chunk_size", 1000),
                    "chunk_overlap": doc.get("chunk_overlap", 100),
                }
                
                record = {
                    "doc_id": doc. get("doc_id", "unknown"),
                    "chunk_id": doc.get("chunk_id", 0),
                    "content": text,
                    "content_type": doc.get("content_type", "text"),
                    "embedding": embedding_list,
                    "metadata":  metadata
                }
                
                self.client.table(self.table_name).insert(record).execute()
                total_added += 1
                
            except Exception as e:
                print(f"❌ Error inserting document: {str(e)}")
        
        print(f"✅ Added {total_added}/{len(documents)} documents to Supabase")
        return total_added
    
    def similarity_search(
        self,
        query: str,
        k: int = 5,
        filter_doc_id: Optional[str] = None,
        score_threshold: float = 0.0
    ) -> List[Dict[str, Any]]:
        """Search for similar documents using pgvector"""
        # Generate query embedding
        query_embedding = self.embed_model.encode([query], convert_to_numpy=True)[0]
        embedding_list = [float(x) for x in query_embedding.tolist()]
        
        print(f"🔍 Searching for:  '{query[: 50]}...'")
        
        try:
            response = self.client.rpc(
                "match_documents",
                {
                    "query_embedding": embedding_list,
                    "match_count": k,
                    "filter_doc_id": filter_doc_id
                }
            ).execute()
            
            print(f"✅ Search returned {len(response. data)} results")
            
            results = []
            for row in response.data:
                similarity = float(row. get("similarity", 0))
                
                if similarity >= score_threshold:
                    metadata = row.get("metadata", {})
                    if isinstance(metadata, str):
                        try:
                            metadata = json. loads(metadata)
                        except:
                            metadata = {}
                    
                    results.append({
                        "id": row.get("id"),
                        "doc_id": row.get("doc_id"),
                        "chunk_id": row.get("chunk_id"),
                        "text": row.get("content"),
                        "content_type": row.get("content_type"),
                        "metadata": metadata,
                        "score": similarity
                    })
                    
                    print(f"  - Score: {similarity:.4f} | {row.get('doc_id')} | {row.get('content', '')[:50]}...")
            
            return results
        
        except Exception as e: 
            print(f"❌ Search error: {str(e)}")
            import traceback
            traceback.print_exc()
            return []
    
    def delete_by_doc_id(self, doc_id: str) -> int:
        """Delete all chunks for a specific document"""
        try: 
            response = self.client. table(self.table_name).delete().eq("doc_id", doc_id).execute()
            deleted_count = len(response.data) if response.data else 0
            print(f"✅ Deleted {deleted_count} chunks for doc_id: {doc_id}")
            return deleted_count
        except Exception as e:
            print(f"❌ Error deleting document: {str(e)}")
            return 0
    
    def delete_all(self) -> int:
        """Delete all documents from the vector store"""
        try: 
            response = self.client.table(self.table_name).delete().neq("id", 0).execute()
            deleted_count = len(response.data) if response.data else 0
            print(f"✅ Deleted all {deleted_count} documents")
            return deleted_count
        except Exception as e:
            print(f"❌ Error deleting all documents: {str(e)}")
            return 0
    
    def get_all_doc_ids(self) -> List[str]:
        """Get all unique document IDs"""
        try:
            response = self.client.table(self. table_name).select("doc_id").execute()
            doc_ids = list(set(row["doc_id"] for row in response.data))
            return doc_ids
        except Exception as e:
            print(f"❌ Error getting doc IDs: {str(e)}")
            return []
    
    def get_document_count(self) -> int:
        """Get total number of chunks in the store"""
        try:
            response = self.client.table(self.table_name).select("id", count="exact").execute()
            return response.count or 0
        except Exception as e:
            print(f"❌ Error getting count: {str(e)}")
            return 0
    
    def document_exists(self, doc_id: str) -> bool:
        """Check if a document already exists"""
        try:
            response = self.client.table(self.table_name).select("id").eq("doc_id", doc_id).limit(1).execute()
            return len(response.data) > 0
        except Exception as e: 
            print(f"❌ Error checking document:  {str(e)}")
            return False


# ========== Reranking Functions ==========
def rerank_results(
    query: str,
    results: List[Dict[str, Any]],
    embed_model: SentenceTransformer,
    rerank_method: str = "hybrid"
) -> List[Dict[str, Any]]:
    """Rerank search results"""
    if not results:
        return results
    
    query_words = set(query.lower().split())
    scored_results = []
    
    for result in results:
        text = result["text"]
        original_score = result["score"]
        
        doc_words = set(text.lower().split())
        keyword_overlap = len(query_words. intersection(doc_words)) / max(1, len(query_words))
        keyword_score = keyword_overlap
        
        if rerank_method == "hybrid":
            combined_score = 0.7 * original_score + 0.3 * keyword_score
        elif rerank_method == "keyword":
            combined_score = keyword_score
        else:
            combined_score = original_score
        
        scored_results.append({
            **result,
            "original_score": float(original_score),
            "keyword_score": float(keyword_score),
            "combined_score": float(combined_score)
        })
    
    scored_results.sort(key=lambda x: x["combined_score"], reverse=True)
    
    return scored_results


def search_vector_store_with_reranking(
    vector_store: SupabaseVectorStore,
    query:  str,
    embed_model: SentenceTransformer,
    k: int = 5,
    rerank_top_k: int = 20,
    rerank_method:  str = "hybrid",
    score_threshold: float = 0.0
) -> List[Dict[str, Any]]: 
    """Search vector store dengan reranking"""
    # Get initial results from Supabase
    initial_results = vector_store.similarity_search(
        query=query,
        k=rerank_top_k,
        score_threshold=score_threshold
    )
    
    # Rerank results
    reranked_results = rerank_results(query, initial_results, embed_model, rerank_method)
    
    return reranked_results[: k]


# ========== Language Detection ==========
def detect_language(text: str) -> str:
    """Simple language detection based on character patterns"""
    indonesian_words = ['yang', 'dan', 'di', 'ke', 'dari', 'untuk', 'pada', 'adalah', 'dalam',
                        'dengan', 'ini', 'itu', 'atau', 'juga', 'akan', 'telah', 'ada',
                        'apa', 'siapa', 'bagaimana', 'mengapa', 'kapan', 'dimana']
    
    english_words = ['the', 'is', 'are', 'was', 'were', 'what', 'who', 'how', 'why', 'when',
                     'where', 'which', 'this', 'that', 'these', 'those', 'have', 'has', 'had']
    
    text_lower = text.lower()
    words = text_lower.split()
    
    indonesian_count = sum(1 for word in words if word in indonesian_words)
    english_count = sum(1 for word in words if word in english_words)
    
    if english_count > indonesian_count:
        return 'english'
    else:
        return 'indonesian'


def create_enhanced_context(chunks: List[DocChunk]) -> str:
    """Create enhanced context from document chunks with page references"""
    parts = []
    for ch in chunks:
        content_type = ch.meta.get('content_type', 'text')
        page_info = f", Halaman {ch.page_number}" if ch.page_number else ""
        header = f"[{ch. doc_id}{page_info} | {ch.meta.get('category', 'unknown')} | Tipe: {content_type}]\n"
        parts.append(header + ch.text. strip())
    return "\n\n---\n\n".join(parts)


def results_to_doc_chunks(results: List[Dict[str, Any]]) -> List[DocChunk]:
    """Convert search results to DocChunk objects"""
    chunks = []
    for result in results:
        metadata = result.get("metadata", {})
        chunk = DocChunk(
            doc_id=result. get("doc_id", "unknown"),
            chunk_id=result.get("chunk_id", 0),
            text=result["text"],
            meta={
                "category": metadata.get("category", "unknown"),
                "keywords": metadata.get("keywords", ""),
                "content_type": result.get("content_type", "text"),
                "doc_id": result.get("doc_id", "unknown"),
                "chunk_id": result.get("chunk_id", 0),
            },
            content_type=result.get("content_type", "text"),
            page_number=metadata.get("page_number")
        )
        chunks.append(chunk)
    return chunks


def answer_with_rag(
    query: str, 
    retrieved: List[DocChunk], 
    chat_model: str,
    chat_history: Optional[List[Dict[str, str]]] = None
) -> str:
    """
    Generate an answer using:
    1. Retrieved document snippets (from vector search)  
    2. Chat history (for conversational context)
    3. User query
    
    Flow:
    - Documents provide the FACTS
    - Chat history provides CONTEXT for follow-up questions
    - LLM synthesizes an answer from both
    """
    client = get_groq_client()
    language = detect_language(query)
    
    # Build document context from retrieved chunks
    doc_context = create_enhanced_context(retrieved)
    
    # Build chat history summary (last 6 messages = 3 exchanges)
    history_context = ""
    if chat_history and len(chat_history) > 0:
        recent_history = chat_history[-6:] if len(chat_history) > 6 else chat_history
        history_parts = []
        for msg in recent_history:
            role_label = "Pengguna" if msg["role"] == "user" else "Asisten"
            # Truncate long messages
            content = msg["content"]
            if len(content) > 400:
                content = content[:400] + "..."
            history_parts.append(f"{role_label}: {content}")
        history_context = "\n".join(history_parts)
    
    if language == "indonesian":
        system_prompt = """Anda adalah asisten AI untuk analisis dokumen yang dapat melakukan percakapan natural.

CARA KERJA ANDA:
1. KONTEKS DOKUMEN berisi kutipan dari dokumen yang diunggah pengguna - ini adalah SUMBER UTAMA jawaban Anda
2. RIWAYAT PERCAKAPAN membantu Anda memahami konteks pertanyaan follow-up
3. Kombinasikan keduanya untuk memberikan jawaban yang relevan dan akurat

CONTOH PENGGUNAAN KONTEKS:
- Jika pengguna sebelumnya bertanya "apa itu order?" dan sekarang bertanya "berikan contohnya"
- Anda harus memberikan contoh tentang ORDER berdasarkan dokumen

ATURAN: 
✓ Jawab berdasarkan DOKUMEN yang disediakan
✓ Gunakan riwayat untuk memahami maksud pertanyaan
✓ Berikan jawaban yang informatif dan terstruktur
✓ Sertakan referensi dokumen
✗ JANGAN mengarang informasi yang tidak ada di dokumen
✗ JANGAN mengabaikan konteks percakapan"""

        user_prompt = f"""📜 RIWAYAT PERCAKAPAN:
{history_context if history_context else "(Percakapan baru)"}

📄 KONTEKS DOKUMEN:
{doc_context}

❓ PERTANYAAN:  {query}

Jawab pertanyaan di atas berdasarkan konteks dokumen.  Pahami maksud pertanyaan dari riwayat percakapan jika diperlukan. 

Format jawaban:
1. Jawaban langsung dan informatif
2. Penjelasan atau contoh jika relevan  
3. Referensi: [nama_dokumen], halaman [nomor]"""

    else:
        system_prompt = """You are an AI assistant for document analysis that can hold natural conversations.

HOW YOU WORK:
1. DOCUMENT CONTEXT contains excerpts from user-uploaded documents - this is your PRIMARY SOURCE
2. CHAT HISTORY helps you understand follow-up questions
3. Combine both to provide relevant and accurate answers

EXAMPLE: 
- If user previously asked "what is order?" and now asks "give an example"
- You should provide examples about ORDER based on the documents

RULES:
✓ Answer based on PROVIDED DOCUMENTS
✓ Use history to understand question intent
✓ Give informative, structured answers
✓ Include document references
✗ DO NOT make up information not in documents
✗ DO NOT ignore conversation context"""

        user_prompt = f"""📜 CHAT HISTORY: 
{history_context if history_context else "(New conversation)"}

📄 DOCUMENT CONTEXT:
{doc_context}

❓ QUESTION:  {query}

Answer the question above based on document context.  Understand the question intent from chat history if needed.

Format: 
1. Direct, informative answer
2. Explanation or examples if relevant
3. References: [document_name], page [number]"""

    try:
        resp = client.chat.completions.create(
            model=chat_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=1500,
        )
        return resp.choices[0].message.content
    except Exception as e: 
        error_msg = str(e)
        if language == "indonesian":
            return f"Kesalahan saat menghasilkan jawaban: {error_msg}"
        else: 
            return f"Error generating answer:  {error_msg}"